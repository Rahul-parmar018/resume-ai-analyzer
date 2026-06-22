"""
Text Extraction Service
Supports PDF and DOCX file formats
"""

import io
from PyPDF2 import PdfReader
from docx import Document


def extract_text_from_pdf(file_bytes: bytes) -> str:
    """
    Extract text from PDF file

    Args:
        file_bytes: Raw PDF file bytes

    Returns:
        Extracted text as string
    """
    try:
        pdf_reader = PdfReader(io.BytesIO(file_bytes))
        text_parts = []

        for page in pdf_reader.pages:
            text = page.extract_text()
            if text:
                text_parts.append(text)

        full_text = "\n".join(text_parts)

        if not full_text.strip():
            raise ValueError("PDF contains no extractable text")

        return full_text

    except Exception as e:
        raise ValueError(f"Failed to extract PDF text: {str(e)}")


def extract_text_from_docx(file_bytes: bytes) -> str:
    """
    Extract text from DOCX file

    Args:
        file_bytes: Raw DOCX file bytes

    Returns:
        Extracted text as string
    """
    try:
        doc = Document(io.BytesIO(file_bytes))
        text_parts = []

        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text_parts.append(paragraph.text)

        # Also extract tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        text_parts.append(cell.text)

        full_text = "\n".join(text_parts)

        if not full_text.strip():
            raise ValueError("DOCX contains no extractable text")

        return full_text

    except Exception as e:
        raise ValueError(f"Failed to extract DOCX text: {str(e)}")


def extract_text(file_bytes: bytes, filename: str) -> dict:
    """
    Main entry point - extracts text based on file extension

    Args:
        file_bytes: Raw file bytes
        filename: Original filename (used to detect file type)

    Returns:
        {
            "text": "extracted content",
            "filename": "original filename",
            "file_type": "pdf" or "docx",
            "char_count": number of characters
        }

    Raises:
        ValueError: If file type is unsupported or extraction fails
    """
    # Detect file type from extension
    ext = filename.lower().split('.')[-1]

    if ext == 'pdf':
        file_type = 'pdf'
        text = extract_text_from_pdf(file_bytes)
    elif ext in ['docx', 'doc']:
        file_type = 'docx'
        text = extract_text_from_docx(file_bytes)
    else:
        raise ValueError(f"Unsupported file type: {ext}. Only PDF and DOCX are supported.")

    # Clean up the text
    text = clean_text(text)

    return {
        "text": text,
        "filename": filename,
        "file_type": file_type,
        "char_count": len(text)
    }


def clean_text(text: str) -> str:
    """
    Clean and normalize extracted text
    """
    # Remove excessive whitespace
    lines = text.split('\n')
    cleaned_lines = [line.strip() for line in lines if line.strip()]
    return '\n'.join(cleaned_lines)